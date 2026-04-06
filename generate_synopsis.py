from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== Page Setup (A4) =====
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== Styles =====
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(51, 51, 51)

# Helper functions
def add_heading_styled(text, level=1, color=RGBColor(10, 15, 30)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_accent_line():
    p = doc.add_paragraph()
    run = p.add_run('━' * 12)
    run.font.color.rgb = RGBColor(220, 38, 38)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(' ')
    run.font.size = Pt(int(pts / 2))

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color='D4D4D4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')
        tcBorders.append(el)
    tcPr.append(tcBorders)


# ============================================================
# PAGE 1 — COVER
# ============================================================
add_spacer(80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROGRAM SYNOPSIS · 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True
run.font.all_caps = True

add_spacer(20)

# Logo reference
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', 'images', 'NUDGE LOGO.png')
if os.path.exists(logo_path):
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(1.0))

add_spacer(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nudge ')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(10, 15, 30)
run = p.add_run('Fellowship')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(220, 38, 38)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COHORT 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(140, 140, 140)
run.font.bold = True

add_spacer(8)

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━━━━━━━━')
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.size = Pt(10)

add_spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A fully-funded, transformative week in New Delhi for Kerala\'s\nbrightest Class 12 students. Explore premier campuses.\nLearn from experts. Return inspired.')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 100, 100)
p.paragraph_format.line_spacing = Pt(22)

add_spacer(20)

# Stats row as table
stats_table = doc.add_table(rows=2, cols=4)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_data = [
    ('6', 'Fellows'), ('3', 'Streams'), ('5+', 'Campuses'), ('₹0', 'Cost')
]
for i, (num, label) in enumerate(stats_data):
    cell = stats_table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(num)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(10, 15, 30)

    cell2 = stats_table.cell(1, i)
    cell2.text = ''
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(label.upper())
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(140, 140, 140)
    run2.font.bold = True

add_spacer(60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('nudge.academy/fellowship · © 2026 Nudge Academy')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(160, 160, 160)

doc.add_page_break()


# ============================================================
# PAGE 2 — PROGRAM OVERVIEW
# ============================================================
p = doc.add_paragraph()
run = p.add_run('PROGRAM OVERVIEW')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('What is the Nudge Fellowship?', level=1)

p = doc.add_paragraph()
run = p.add_run('The Nudge Fellowship is a fully-funded, immersive program that takes 6 exceptional Class 12 students from Kerala on a transformative week-long journey to New Delhi. Fellows explore premier university campuses, attend expert-led sessions, experience India\'s rich heritage, and build a lifelong peer network — all at zero cost to the student.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 80, 80)
p.paragraph_format.line_spacing = Pt(20)

add_spacer(8)

# Highlights as table (2x2 grid)
hl_data = [
    ('🎯 Our Mission', 'Help students see their future firsthand — walk through the campuses, meet the people, and experience the life they\'re working toward.'),
    ('⚖️ Equal Representation', '2 fellows from Science, 2 from Commerce, and 2 from Humanities — ensuring every stream has a seat at the table.'),
    ('💰 100% Funded', 'Train travel, accommodation, meals, campus transport, and all activities are fully covered by Nudge Academy.'),
    ('📍 New Delhi', 'India\'s capital — home to the country\'s most prestigious universities, iconic monuments, and vibrant culture.'),
]

hl_table = doc.add_table(rows=2, cols=2)
hl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (title, desc) in enumerate(hl_data):
    row = i // 2
    col = i % 2
    cell = hl_table.cell(row, col)
    cell.text = ''
    set_cell_shading(cell, 'F8F9FA')
    set_cell_borders(cell, 'E8E8E8')
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(10, 15, 30)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = RGBColor(100, 100, 100)
    p2.paragraph_format.space_before = Pt(4)

add_spacer(12)
add_accent_line()

p = doc.add_paragraph()
run = p.add_run('WHAT\'S COVERED')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('Zero Cost to the Fellow', level=2)

covered_table = doc.add_table(rows=1, cols=4)
covered_table.alignment = WD_TABLE_ALIGNMENT.CENTER
covered_data = [
    ('🚂', 'Train Travel', 'Round-trip tickets fully covered'),
    ('🏨', 'Accommodation', 'Safe, comfortable stay in Delhi'),
    ('🍽️', 'All Meals', 'Breakfast, lunch & dinner included'),
    ('🚌', 'Local Transport', 'All campus & city travel arranged'),
]
for i, (emoji, title, desc) in enumerate(covered_data):
    cell = covered_table.cell(0, i)
    cell.text = ''
    set_cell_shading(cell, 'FFF5F5')
    set_cell_borders(cell, 'FECACA')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(emoji)
    run.font.size = Pt(20)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(10)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(10, 15, 30)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(desc)
    run3.font.size = Pt(8.5)
    run3.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()


# ============================================================
# PAGE 3 — ELIGIBILITY + SELECTION CRITERIA
# ============================================================
p = doc.add_paragraph()
run = p.add_run('ELIGIBILITY')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('Who Can Apply?', level=1)

p = doc.add_paragraph()
run = p.add_run('The fellowship is open to ambitious students who meet the following criteria:')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

add_spacer(6)

elig_items = [
    ('Kerala Students Only', 'Must be a current resident studying in a recognized Kerala school.'),
    ('Completed Class 11', 'Must have completed Class 11 in Science, Commerce, or Humanities stream.'),
    ('Open to All Streams', 'Students from any stream can apply. Equal representation will be given to all streams — 2 fellows per stream.'),
]

for title, desc in elig_items:
    p = doc.add_paragraph()
    run = p.add_run('● ')
    run.font.color.rgb = RGBColor(220, 38, 38)
    run.font.size = Pt(11)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(10, 15, 30)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.7)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = RGBColor(100, 100, 100)
    add_spacer(4)

add_accent_line()

p = doc.add_paragraph()
run = p.add_run('SELECTION CRITERIA')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('What We\'re Looking For', level=2)

p = doc.add_paragraph()
run = p.add_run('We don\'t just seek achievers — we seek thinkers and doers ready to shape the future.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

add_spacer(6)

values_data = [
    ('💡 Intellectual Curiosity', 'You question, explore, and connect dots across disciplines beyond textbooks.'),
    ('🌏 Engagement with the World', 'A deep awareness of social issues and a drive to understand the world around you.'),
    ('🧭 Drive & Self-Awareness', 'You know your strengths, take initiative, and create your own opportunities.'),
    ('👥 Leadership Potential', 'You lead through empathy and influence, inspiring others toward a common goal.'),
]

val_table = doc.add_table(rows=2, cols=2)
val_table.alignment = WD_TABLE_ALIGNMENT.CENTER
colors = ['FEF2F2', 'FFFBEB', 'EFF6FF', 'ECFDF5']
border_colors = ['FECACA', 'FDE68A', 'BFDBFE', 'A7F3D0']

for i, (title, desc) in enumerate(values_data):
    row = i // 2
    col = i % 2
    cell = val_table.cell(row, col)
    cell.text = ''
    set_cell_shading(cell, colors[i])
    set_cell_borders(cell, border_colors[i])
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(10, 15, 30)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = RGBColor(100, 100, 100)
    p2.paragraph_format.space_before = Pt(4)

doc.add_page_break()


# ============================================================
# PAGE 4 — THE EXPERIENCE
# ============================================================
p = doc.add_paragraph()
run = p.add_run('THE EXPERIENCE')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('What Fellows Will Experience', level=1)

p = doc.add_paragraph()
run = p.add_run('A week that redefines how students see education, careers, and the world around them.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

add_spacer(6)

exp_data = [
    ('🏛️', 'Campus Visits', 'Experience premier campus life at IIT Delhi, JNU, AMU, JMI, and DU.'),
    ('👨‍🏫', 'Expert Sessions', 'Learn from professors and industry leaders who challenge your thinking.'),
    ('🕌', 'Heritage & Culture', 'Explore Delhi\'s rich history, iconic monuments, and vibrant culture.'),
    ('🤝', 'Peer Network', 'Bond with 5 exceptional peers from across Kerala. Build a lifelong network.'),
    ('🧠', 'Self-Discovery', 'Uncover your unique strengths through workshops and guided reflections.'),
    ('📜', 'Certificate & Report', 'Receive a Fellowship certificate and a personal outcomes report.'),
]

exp_table = doc.add_table(rows=2, cols=3)
exp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (emoji, title, desc) in enumerate(exp_data):
    row = i // 3
    col = i % 3
    cell = exp_table.cell(row, col)
    cell.text = ''
    set_cell_shading(cell, 'F8F9FA')
    set_cell_borders(cell, 'E5E7EB')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(emoji)
    run.font.size = Pt(22)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(10, 15, 30)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(desc)
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(100, 100, 100)
    p3.paragraph_format.space_before = Pt(3)

add_spacer(12)
add_accent_line()

p = doc.add_paragraph()
run = p.add_run('CAMPUS DESTINATIONS')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('Where We\'ll Go', level=2)

campus_data = [
    ('IIT Delhi', 'Indian Institute of Technology'),
    ('JNU', 'Jawaharlal Nehru University'),
    ('AMU', 'Aligarh Muslim University'),
    ('JMI', 'Jamia Millia Islamia'),
    ('Delhi University', 'University of Delhi'),
    ('& More', 'Additional surprise visits'),
]

campus_table = doc.add_table(rows=2, cols=3)
campus_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (name, full) in enumerate(campus_data):
    row = i // 3
    col = i % 3
    cell = campus_table.cell(row, col)
    cell.text = ''
    set_cell_shading(cell, 'F8F9FA')
    set_cell_borders(cell, 'E5E7EB')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(10, 15, 30)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(full)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()


# ============================================================
# PAGE 5 — TIMELINE + SUMMARY
# ============================================================
p = doc.add_paragraph()
run = p.add_run('TIMELINE')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('How It Works', level=1)

add_spacer(4)

timeline_data = [
    ('Till April 25, 2026', 'Applications Open', 'Fill out the online application form at nudge.academy/fellowship — takes approximately 10 minutes.'),
    ('Early May 2026', 'Results Announced', '6 fellows selected — 2 from each stream (Science, Commerce, Humanities) — and notified via email & phone.'),
    ('Mid May 2026', 'Board the Train to Delhi', 'Your fully-funded fellowship week begins. Travel from Kerala to New Delhi with your cohort.'),
    ('Late May 2026', 'Return Transformed', 'Return home with new skills, perspectives, a strong peer network, and a Fellowship certificate.'),
]

for i, (date, title, desc) in enumerate(timeline_data):
    p = doc.add_paragraph()
    run = p.add_run(f'  ●  ')
    run.font.color.rgb = RGBColor(220, 38, 38)
    run.font.size = Pt(10)

    run = p.add_run(date.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.0)
    run2 = p2.add_run(title)
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(10, 15, 30)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(1.0)
    run3 = p3.add_run(desc)
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(100, 100, 100)
    p3.paragraph_format.space_after = Pt(12)

add_accent_line()

p = doc.add_paragraph()
run = p.add_run('AT A GLANCE')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_heading_styled('Fellowship Summary', level=2)

summary_data = [
    ('Program', 'Nudge Fellowship 2026'),
    ('Duration', '1 Week (approximately 7 days)'),
    ('Location', 'New Delhi, India'),
    ('Fellows', '6 students — 2 from each stream'),
    ('Streams', 'Science, Commerce, Humanities'),
    ('Deadline', 'April 25, 2026'),
    ('Cost', '₹0 — 100% Fully Funded'),
    ('Target Group', 'Class 12 students from Kerala'),
    ('Organized By', 'Nudge Academy'),
]

sum_table = doc.add_table(rows=len(summary_data), cols=2)
sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(summary_data):
    cell1 = sum_table.cell(i, 0)
    cell1.text = ''
    set_cell_shading(cell1, 'F8F9FA')
    set_cell_borders(cell1, 'E5E7EB')
    p = cell1.paragraphs[0]
    run = p.add_run(label.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(120, 120, 120)

    cell2 = sum_table.cell(i, 1)
    cell2.text = ''
    set_cell_borders(cell2, 'E5E7EB')
    p2 = cell2.paragraphs[0]
    run2 = p2.add_run(value)
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(10, 15, 30)

doc.add_page_break()


# ============================================================
# PAGE 6 — APPLY + CONTACT
# ============================================================
add_spacer(60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('APPLY NOW')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Ready to See Your Future?')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(10, 15, 30)

add_spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('6 seats. 3 streams. 100% funded.\nApply now and take the first step toward a transformative\nexperience that will change the way you see your future.')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)
p.paragraph_format.line_spacing = Pt(20)

add_spacer(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('nudge.academy/fellowship')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(220, 38, 38)

add_spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('APPLICATION DEADLINE: APRIL 25, 2026')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(140, 140, 140)

add_spacer(40)

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━━━━━━━━━━━━━━━━━━━━')
run.font.color.rgb = RGBColor(230, 230, 230)
run.font.size = Pt(8)

add_spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GET IN TOUCH')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(220, 38, 38)
run.font.bold = True

add_spacer(8)

contact_table = doc.add_table(rows=1, cols=3)
contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
contact_data = [
    ('WEBSITE', 'nudge.academy'),
    ('INSTAGRAM', '@nudge.academy'),
    ('LINKEDIN', 'Nudge Academy'),
]
for i, (label, value) in enumerate(contact_data):
    cell = contact_table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(140, 140, 140)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(value)
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(220, 38, 38)

add_spacer(40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('© 2026 Nudge Academy. All rights reserved.\nBuilt with ❤️ for the dreamers of Kerala.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(170, 170, 170)


# ===== SAVE =====
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Nudge_Fellowship_2026_Synopsis.docx')
doc.save(output_path)
print(f'✅ Word document saved to: {output_path}')
